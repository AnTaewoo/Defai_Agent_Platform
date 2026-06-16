"""문서 포맷 렌더러(md/docx)."""
from __future__ import annotations


def render_md(title: str, body: str, security_level: int, sources: list[str]) -> str:
    """마크다운 문서 렌더. 보안등급 워터마크 + 출처 섹션 포함."""
    watermark = f"// 등급: L{security_level} //"
    lines = [watermark, "", f"# {title}", "", body, ""]
    lines.append("## 출처")
    if sources:
        for source in sources:
            lines.append(f"- {source}")
    else:
        lines.append("- (없음)")
    lines.append("")
    lines.append(watermark)
    return "\n".join(lines)


def render_docx(path: str, title: str, body: str, security_level: int, sources: list[str]) -> None:
    """docx 파일 렌더. 보안등급 워터마크 + 출처 목록 포함."""
    from docx import Document

    watermark = f"// 등급: L{security_level} //"

    document = Document()
    document.add_paragraph(watermark)
    document.add_heading(title, level=1)
    for para in body.split("\n"):
        document.add_paragraph(para)

    document.add_heading("출처", level=2)
    if sources:
        for source in sources:
            document.add_paragraph(source, style="List Bullet")
    else:
        document.add_paragraph("(없음)")

    document.add_paragraph(watermark)
    document.save(path)
