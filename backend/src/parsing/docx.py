"""[P1] Word(.docx) 파서. python-docx로 직접 제어 — LangChain 로더 위임 금지."""

from __future__ import annotations

from docx import Document

from ..types import ParsedDocument


def parse_docx(path: str) -> ParsedDocument:
    """docx 파일을 ParsedDocument로 변환.

    빈 문단은 건너뛰고, 텍스트가 있는 문단을 하나씩 text_blocks에 담아 의미 단위를
    보존한다. 문서 내 표는 각각 tables에 행렬(list[list[str]])로 보존하고,
    빈 셀은 ""로 채운다.
    """
    document = Document(path)

    text_blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_blocks.append(text)

    tables: list[list[list[str]]] = []
    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = ["" if cell.text is None else cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append(rows)

    return ParsedDocument(
        source=path,
        doc_type="docx",
        text_blocks=text_blocks,
        tables=tables,
    )
