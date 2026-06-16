from docx import Document

from src.parsing import parse
from src.parsing.docx import parse_docx


def _make_docx(path: str) -> str:
    document = Document()
    document.add_paragraph("Hello D.A.P")
    document.add_paragraph("")  # 빈 문단 - 건너뛰어야 함
    document.add_paragraph("두 번째 문단")

    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "이름"
    table.cell(0, 1).text = "부서"
    table.cell(0, 2).text = "등급"
    table.cell(1, 0).text = "홍길동"
    table.cell(1, 1).text = "기획"
    table.cell(1, 2).text = "2"

    document.save(path)
    return path


def test_parse_docx_extracts_paragraphs_and_table(tmp_path):
    path = _make_docx(str(tmp_path / "sample.docx"))

    doc = parse_docx(path)

    assert doc.doc_type == "docx"
    assert doc.source == path
    assert "Hello D.A.P" in doc.text_blocks
    assert "두 번째 문단" in doc.text_blocks
    assert "" not in doc.text_blocks
    assert doc.tables == [
        [["이름", "부서", "등급"], ["홍길동", "기획", "2"]]
    ]


def test_parse_dispatches_docx_by_extension(tmp_path):
    path = _make_docx(str(tmp_path / "sample.docx"))

    doc = parse(path)

    assert doc.doc_type == "docx"
